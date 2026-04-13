"""Extração de conteúdo textual de arquivos de upload.

Tipos suportados:
    text:   .md, .txt, .csv
    docs:   .docx, .pdf
    sheets: .xlsx
    image:  .png, .jpg, .jpeg, .webp, .gif (sem extração textual)

A função `extract(path, mime)` retorna (kind, text) onde:
    kind: 'text' | 'image'
    text: string (pode ser vazia para imagens)
"""
from __future__ import annotations
from pathlib import Path
import csv
import io
import mimetypes


IMAGE_MIMES = {"image/png", "image/jpeg", "image/webp", "image/gif", "image/svg+xml"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".svg"}

MAX_CHARS_PER_FILE = 12_000  # limite ao alimentar o prompt


def guess_mime(filename: str) -> str:
    mime, _ = mimetypes.guess_type(filename)
    if mime:
        return mime
    ext = Path(filename).suffix.lower()
    if ext == ".md":
        return "text/markdown"
    if ext == ".csv":
        return "text/csv"
    if ext == ".xlsx":
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if ext == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return "application/octet-stream"


def is_image(mime: str, filename: str = "") -> bool:
    if mime in IMAGE_MIMES or mime.startswith("image/"):
        return True
    return Path(filename).suffix.lower() in IMAGE_EXTS


def _truncate(text: str) -> str:
    if len(text) <= MAX_CHARS_PER_FILE:
        return text
    return text[:MAX_CHARS_PER_FILE] + "\n[...truncado...]"


def _extract_txt(path: Path) -> str:
    for enc in ("utf-8", "latin-1"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_bytes().decode("utf-8", errors="replace")


def _extract_csv(path: Path) -> str:
    out = io.StringIO()
    try:
        with path.open("r", encoding="utf-8", newline="") as f:
            sample = f.read(2048)
            f.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
            except csv.Error:
                dialect = csv.excel
            reader = csv.reader(f, dialect)
            for i, row in enumerate(reader):
                out.write(" | ".join(row) + "\n")
                if i >= 200:
                    out.write("[...demais linhas omitidas...]\n")
                    break
    except Exception as e:
        return f"[erro lendo CSV: {e}]"
    return out.getvalue()


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "[pypdf não instalado]"
    try:
        reader = PdfReader(str(path))
        parts: list[str] = []
        for i, page in enumerate(reader.pages[:50]):
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
            if sum(len(p) for p in parts) > MAX_CHARS_PER_FILE:
                break
        return "\n\n".join(p for p in parts if p.strip())
    except Exception as e:
        return f"[erro lendo PDF: {e}]"


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        return "[python-docx não instalado]"
    try:
        doc = Document(str(path))
        lines: list[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                lines.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    lines.append(" | ".join(cells))
        return "\n".join(lines)
    except Exception as e:
        return f"[erro lendo DOCX: {e}]"


def _extract_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "[openpyxl não instalado]"
    try:
        wb = load_workbook(str(path), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet in wb.sheetnames[:5]:
            ws = wb[sheet]
            parts.append(f"## Planilha: {sheet}")
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                cells = ["" if v is None else str(v) for v in row]
                parts.append(" | ".join(cells))
                if i >= 200:
                    parts.append("[...demais linhas omitidas...]")
                    break
            parts.append("")
        wb.close()
        return "\n".join(parts)
    except Exception as e:
        return f"[erro lendo XLSX: {e}]"


def extract(path: Path, mime: str) -> tuple[str, str]:
    """Retorna (kind, extracted_text). kind ∈ {'text','image'}."""
    filename = path.name
    if is_image(mime, filename):
        return "image", ""

    ext = path.suffix.lower()

    if ext in (".md", ".txt") or mime.startswith("text/plain") or mime == "text/markdown":
        return "text", _truncate(_extract_txt(path))
    if ext == ".csv" or mime == "text/csv":
        return "text", _truncate(_extract_csv(path))
    if ext == ".pdf" or mime == "application/pdf":
        return "text", _truncate(_extract_pdf(path))
    if ext == ".docx" or "wordprocessingml" in mime:
        return "text", _truncate(_extract_docx(path))
    if ext == ".xlsx" or "spreadsheetml" in mime:
        return "text", _truncate(_extract_xlsx(path))
    # fallback: tenta ler como texto
    try:
        return "text", _truncate(_extract_txt(path))
    except Exception:
        return "text", ""
